import random
from datetime import datetime
import sys
from PyQt5.QtCore import Qt
from PyQt5.QtWidgets import QApplication, QMainWindow, QVBoxLayout, QPushButton, QFrame, QLabel, QWidget, \
    QStackedWidget, QHBoxLayout, QSizePolicy, QSpacerItem, QTableWidget, QTableWidgetItem, QHeaderView, QLineEdit, \
    QComboBox
from PyQt5.QtGui import QIcon, QPixmap
from qt_material import apply_stylesheet
from db import connect_to_db
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

class ComputerStoreApp(QMainWindow):
    def __init__(self):
        super(ComputerStoreApp, self).__init__()

        # Подключение к базе данных PostgreSQL
        self.connection = connect_to_db()

        self.setWindowTitle("Магазин компьютерной техники")
        self.setWindowIcon(QIcon("icons/database.png"))

        # Создание боковой панели и области справа
        central_widget = QWidget(self)
        main_layout = QHBoxLayout(central_widget)

        # Боковая панель
        side_panel = QFrame()
        side_panel.setObjectName("side_panel")
        side_panel_layout = QVBoxLayout(side_panel)

        # Добавление изображения в боковую панель
        image_label = QLabel()
        pixmap = QPixmap("icons/compshop.png")  # Загрузка изображения
        image_label.setPixmap(pixmap)
        side_panel_layout.addWidget(image_label)

        # Создание списка для хранения кнопок
        self.buttons = []

        # Добавление кнопок на боковую панель
        buttons_data = [
            ("Заказы", "icons/order.png"),
            ("История заказов", "icons/history_data.png"),
            ("Товары", "icons/order.png"),
            ("Сотрудники", "icons/staff.png"),
            ("Клиенты", "icons/crowd.png"),
            ("Отчет", "icons/report.png"),
        ]

        for idx, (button_text, icon_path) in enumerate(buttons_data):
            button = QPushButton(button_text)
            button.setIcon(QIcon(icon_path))

            # Центрирование текста с .png по левому краю
            button.setStyleSheet("text-align: left; padding-left: 20px;")
            if button_text == "Товары":
                button.clicked.connect(self.show_products_table_page)
                side_panel_layout.addWidget(button)
                self.buttons.append(button)
            elif button_text == "Сотрудники":
                button.clicked.connect(self.show_employees_table_page)
                side_panel_layout.addWidget(button)
                self.buttons.append(button)
            elif button_text == "Клиенты":
                button.clicked.connect(self.show_clients_table_page)
                side_panel_layout.addWidget(button)
                self.buttons.append(button)
            elif button_text == "История заказов":
                button.clicked.connect(self.show_history_orders)
                side_panel_layout.addWidget(button)
                self.buttons.append(button)
            elif button_text == "Отчет":
                button.clicked.connect(self.show_report_page)
                side_panel_layout.addWidget(button)
                self.buttons.append(button)

            else:
                button.clicked.connect(lambda checked, index=idx: self.show_page(index))
                side_panel_layout.addWidget(button)
                self.buttons.append(button)

        # Добавление растягивающегося пространства с высшим приоритетом
        spacer = QSpacerItem(20, 40, QSizePolicy.Minimum, QSizePolicy.Expanding)
        side_panel_layout.addItem(spacer)

        # Область справа
        self.stacked_widget = QStackedWidget()
        main_layout.addWidget(side_panel)
        main_layout.addWidget(self.stacked_widget)

        # Применение стиля Qt Material
        apply_stylesheet(self, theme="light_blue.xml")

        # Установка центрального виджета
        self.setCentralWidget(central_widget)

        # Начальная страница
        self.show_page(0)

    def show_page(self, index):
        if index == 0:  # Страница "Заказы"
            # Удаление кнопки "Просмотр заказа", если она присутствует
            for i in reversed(range(self.stacked_widget.count())):
                page_widget = self.stacked_widget.widget(i)
                if isinstance(page_widget, QWidget):
                    for child in page_widget.children():
                        if isinstance(child, QPushButton) and child.text() == "Просмотр заказа":
                            child.deleteLater()

            page_widget = QWidget()
            page_layout = QVBoxLayout(page_widget)

            # Добавление надписи "Заказы", кнопки "Показ таблицы" и кнопки "Просмотр заказа" на одной горизонтали
            page_title_layout = QHBoxLayout()
            label_title = QLabel("Заказы")
            label_title.setAlignment(Qt.AlignCenter)
            label_title.setStyleSheet("font-size: 20px; font-weight: bold;")  # Установка размера и жирного шрифта
            page_title_layout.addWidget(label_title)
            show_table_button = QPushButton("Показ таблицы")
            show_table_button.clicked.connect(
                lambda: self.show_page(index))  # Переключение на текущую страницу (Заказы)
            show_table_button.setStyleSheet("background-color: transparent; border: none;")
            page_title_layout.addWidget(show_table_button)
            # view_order_button = QPushButton("Просмотр заказа")
            # view_order_button.clicked.connect(lambda: self.show_order_details())
            # view_order_button.setStyleSheet("background-color: transparent; border: none;")
            # page_title_layout.addWidget(view_order_button)
            # page_title_layout.addStretch(1)  # Добавление растягивающегося пространства для разделения элементов
            page_layout.addLayout(page_title_layout)

            # Создание таблицы для отображения заказов
            table_widget = QTableWidget()
            table_widget.setColumnCount(4)
            table_widget.setHorizontalHeaderLabels(["Код заказа", "Номер заказа", "ФИО пользователя", "Действия"])
            table_widget.verticalHeader().setVisible(False)

            # Добавление данных из базы данных
            cursor = self.connection.cursor()
            cursor.execute("SELECT * FROM Заказы;")
            rows = cursor.fetchall()

            for row_num, row_data in enumerate(rows):
                table_widget.insertRow(row_num)
                for column_num, data in enumerate(row_data):
                    item = QTableWidgetItem(str(data))
                    if column_num < 3:
                        item.setFlags(
                            item.flags() & ~Qt.ItemIsEditable)  # Сделать поля кода, номера заказа и ФИО нередактируемыми
                    table_widget.setItem(row_num, column_num, item)

                # Создание кнопок действий
                action_widget = QWidget()
                action_layout = QHBoxLayout(action_widget)
                view_button = QPushButton()
                view_button.setIcon(QIcon("icons/eye.png"))
                view_button.setStyleSheet("border: none;")
                view_button.clicked.connect(lambda checked, row=row_data: self.show_product_details(row[0]))
                confirm_button = QPushButton()
                confirm_button.setIcon(QIcon("icons/done.png"))
                confirm_button.setStyleSheet("border: none;")
                confirm_button.clicked.connect(lambda checked, row=row_data: self.confirm_order(row_data[0]))
                delete_button = QPushButton()
                delete_button.setIcon(QIcon("icons/delete.png"))
                delete_button.setStyleSheet("border: none;")
                delete_button.clicked.connect(lambda checked, row=row_data: self.delete_order(row_data[0]))
                action_layout.addWidget(view_button)
                action_layout.addWidget(confirm_button)
                action_layout.addWidget(delete_button)
                action_layout.setContentsMargins(0, 0, 0, 0)
                action_layout.setSpacing(0)
                action_widget.setLayout(action_layout)
                table_widget.setCellWidget(row_num, 3, action_widget)

            cursor.close()

            # Растягиваем таблицу по горизонтали и вертикали
            table_widget.horizontalHeader().setStretchLastSection(True)
            table_widget.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
            table_widget.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)

            page_layout.addWidget(table_widget)
            self.stacked_widget.addWidget(page_widget)
            self.stacked_widget.setCurrentWidget(page_widget)

        else:
            # В остальных случаях пока что создаем заглушку
            page_widget = QWidget()
            label = QLabel(f"Функционал для страницы '{self.buttons_data[index][0]}'")
            label.setAlignment(Qt.AlignCenter)
            page_layout = QVBoxLayout(page_widget)
            page_layout.addWidget(label)
            self.stacked_widget.addWidget(page_widget)
            self.stacked_widget.setCurrentIndex(index)

        # Отмена выделения предыдущей кнопки
        for button in self.buttons:
            button.setStyleSheet("text-align: left; padding-left: 20px;")

        # Выделение текущей кнопки жирным шрифтом
        self.buttons[index].setStyleSheet("text-align: left; padding-left: 20px; font-weight: bold;")

    def confirm_order(self, order_id):
        # Получение данных подтвержденного заказа
        cursor = self.connection.cursor()
        cursor.execute(
            f"SELECT Номер_заказа, ФИО_клиента, Код_товара, Количество, Цена, Общая_цена FROM Заказы WHERE Код_заказа = {order_id};")
        order_data = cursor.fetchone()

        # Проверка на наличие данных заказа
        if order_data is None:
            print(f"Заказа с кодом {order_id} не существует!")
            return

        # Генерируем случайный 'Код_заказа'
        new_order_id = random.randint(0, 1000000)

        # Перенос заказа в историю с добавлением нового 'Код_заказа', изменением статуса на "Оплачено" и сегодняшней датой
        cursor.execute(f"""
            INSERT INTO История_заказов(Код_заказа, Номер_заказа, ФИО_клиента, Статус, Дата, Код_товара, Количество, Цена, Общая_цена)
            VALUES ({new_order_id}, '{order_data[0]}', '{order_data[1]}', 'Оплачено', '{datetime.now().date().strftime("%Y-%m-%d")}',
            '{order_data[2]}','{order_data[3]}','{order_data[4]}', '{order_data[5]}')
        """)

        # Фактическое удаление заказа
        cursor.execute(f"DELETE FROM Заказы WHERE Код_заказа = {order_id};")

        # Применение изменений
        self.connection.commit()
        cursor.close()

        # Оставаться на той же странице после выполнения
        self.show_products_table_page()

    def show_product_details(self, order_id):
        # Создание новой таблицы с информацией о заказе
        detail_widget = QWidget()
        detail_layout = QVBoxLayout(detail_widget)

        # Добавление заголовка и кнопок на одной горизонтали
        page_title_layout = QHBoxLayout()
        label_title = QLabel("Заказы")
        label_title.setAlignment(Qt.AlignCenter)
        label_title.setStyleSheet("font-size: 20px; font-weight: bold;")
        page_title_layout.addWidget(label_title)

        show_table_button = QPushButton("Показ таблицы")
        show_table_button.clicked.connect(self.show_page)
        show_table_button.setStyleSheet("background-color: transparent; border: none;")
        page_title_layout.addWidget(show_table_button)

        detail_layout.addLayout(page_title_layout)

        detail_table = QTableWidget()
        detail_table.setColumnCount(4)
        detail_table.setHorizontalHeaderLabels(
            ["Наименование товара", "Кол-во", "Цена", "Общая цена"])

        # Установить политику размера виджета
        detail_table.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)

        # Установить режим растягивания для заголовков таблицы
        detail_table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)

        cursor = self.connection.cursor()
        cursor.execute(f"SELECT Код_товара, Количество,Цена, Общая_цена FROM Заказы WHERE Код_заказа = {order_id};")
        rows = cursor.fetchall()
        print(rows)
        for row_num, row_data in enumerate(rows):
            detail_table.insertRow(row_num)

            # Получение имени товара
            product_id = row_data[0]
            print("Впфмц", product_id)
            cursor.execute(f"SELECT Наименование FROM Товары WHERE Код_товара = {product_id};")
            product_name = cursor.fetchone()

            if product_name is not None:
                product_name = product_name[0]
            else:
                product_name = "Неопределенный товар"

            # Замена кода товара на наименование товара
            row_data = list(row_data)
            row_data[0] = product_name

            for column_num, data in enumerate(row_data):
                item = QTableWidgetItem(str(data))
                detail_table.setItem(row_num, column_num, item)

        detail_layout.addWidget(detail_table)
        self.stacked_widget.addWidget(detail_widget)
        self.stacked_widget.setCurrentWidget(detail_widget)

    def show_order_details(self, order_data=None):
        # Создание виджета для отображения данных заказа
        order_widget = QWidget()
        order_layout = QVBoxLayout(order_widget)

        # Добавление надписи "Заказы", кнопки "Показ таблицы" и кнопки "Просмотр заказа" на одной горизонтальной линии
        page_title_layout = QHBoxLayout()
        label_title = QLabel("Заказы")
        label_title.setAlignment(Qt.AlignCenter)
        label_title.setStyleSheet("font-size: 20px; font-weight: bold;")  # Установка размера и жирного шрифта
        page_title_layout.addWidget(label_title)
        show_table_button = QPushButton("Показ таблицы")
        show_table_button.clicked.connect(
            lambda: self.show_page(0))  # Переключение на текущую страницу (Заказы)
        show_table_button.setStyleSheet("background-color: transparent; border: none;")
        page_title_layout.addWidget(show_table_button)
        page_title_layout.addStretch(1)  # Добавление растягивающегося пространства для разделения элементов
        order_layout.addLayout(page_title_layout)

        if order_data:
            # Добавление данных заказа
            order_info_layout = QVBoxLayout()

            for key, value in zip(["Код заказа", "Номер заказа", "ФИО пользователя"], order_data):
                label = QLabel(f"{key}: {value}")
                order_info_layout.addWidget(label)

            order_layout.addLayout(order_info_layout)

        self.stacked_widget.addWidget(order_widget)
        self.stacked_widget.setCurrentWidget(order_widget)

    def show_products_page(self):
        # Удаление всех виджетов из QStackedWidget
        while self.stacked_widget.count() > 0:
            widget = self.stacked_widget.widget(0)
            self.stacked_widget.removeWidget(widget)
            widget.deleteLater()

        products_page_widget = QWidget()
        main_layout = QVBoxLayout(products_page_widget)

        # Добавление надписи "Заказы" и кнопок "Показ таблицы" и "Просмотр заказа"
        title_layout = QHBoxLayout()
        title_label = QLabel("Товары")
        title_label.setStyleSheet("font-size: 20px; font-weight: bold;")
        title_label.setAlignment(Qt.AlignCenter)
        show_table_button = QPushButton("Показ таблицы")
        show_table_button.clicked.connect(self.show_products_table_page)
        show_table_button.setStyleSheet("border: none;")

        add_tovar_button = QPushButton("Добавление товара")
        add_tovar_button.setStyleSheet("background-color: transparent; border: none;")

        title_layout.addWidget(title_label, alignment=Qt.AlignCenter)
        title_layout.addWidget(show_table_button)
        title_layout.addWidget(add_tovar_button )
        title_layout.addStretch()
        main_layout.addLayout(title_layout)

        # Создание разметки для полей ввода
        fields_layout = QVBoxLayout()
        fields = ['Наименование', 'Категория', 'Описание', 'Цена', 'Количество', 'Код товара']
        self.product_fields = {}
        for field in fields:
            field_layout = QVBoxLayout()
            label = QLabel(field)
            line_edit = QLineEdit()
            line_edit.setFixedWidth(200)  # Уменьшение ширины поля для ввода
            self.product_fields[field] = line_edit
            field_layout.addWidget(label)
            field_layout.addWidget(line_edit)
            fields_layout.addLayout(field_layout)

        # Область для добавления изображения
        image_layout = QVBoxLayout()
        image_layout.setSpacing(0)
        image_label = QLabel("Изображение")
        image_line_edit = QLineEdit()
        image_line_edit.setFixedHeight(100)  # Увеличиваем высоту поля
        image_line_edit.setFixedWidth(200)  # Уменьшение ширины поля для ввода изображения
        self.product_fields['Изображение'] = image_line_edit
        image_layout.addWidget(image_label, alignment=Qt.AlignTop)  # Расположить сверху
        image_layout.addWidget(image_line_edit, alignment=Qt.AlignTop)  # Расположить сверху

        # Горизонтальная разметка для объединения полей и загрузки изображения
        content_layout = QHBoxLayout()
        content_layout.addLayout(fields_layout)  # Добавление макета с полями ввода
        content_layout.addLayout(image_layout)  # Добавление макета загрузки изображения справа

        main_layout.addLayout(content_layout)

        # Кнопка подтверждения, расположенная по центру
        confirm_button_layout = QHBoxLayout()
        confirm_button = QPushButton('Подтверждение')
        confirm_button .setStyleSheet("background-color: black; color: white; border: none; padding: 10px;")
        confirm_button.setSizePolicy(QSizePolicy.Maximum, QSizePolicy.Preferred)
        confirm_button_layout.addStretch()  # Выравнивание слева от кнопки
        confirm_button_layout.addWidget(confirm_button)
        confirm_button_layout.addStretch()  # Выравнивание справа от кнопки
        confirm_button.clicked.connect(self.on_confirm_clicked)
        main_layout.addLayout(confirm_button_layout)

        products_page_widget.setLayout(main_layout)
        self.stacked_widget.addWidget(products_page_widget)
        self.stacked_widget.setCurrentWidget(products_page_widget)

    def on_confirm_clicked(self):
        # Сбор данных из полей ввода
        product_data = {key: line_edit.text().strip() for key, line_edit in self.product_fields.items() if
                        key != 'Изображение'}

        # Проверяем обязательные поля на заполненность
        if any(not value for value in product_data.values()):
            print("Не все обязательные поля заполнены!")
            return
        try:
            cursor = self.connection.cursor()
            # Формируем кортеж данных
            data_tuple = tuple(product_data[field] for field in
                               ['Код товара', 'Наименование', 'Категория', 'Описание', 'Цена', 'Количество'])
            cursor.execute(
                "INSERT INTO Товары (Код_товара, Наименование, Категория, Описание, Цена, Количество) VALUES (%s, %s, %s, %s, %s, %s)",
                data_tuple)
            self.connection.commit()  # Не забудьте подтвердить изменения в БД
            print("Данные успешно добавлены в базу.")
            self.show_products_table_page()
        except Exception as e:
            print(f"Произошла ошибка во время добавления данных в БД: {e}")

        finally:
            cursor.close()  # Важно закрыть курсор
    def edit_product(self, product_data):
        # Заполните поля формы данными продукта
        for field_name, line_edit in self.product_fields.items():
            value = str(product_data[field_name])
            line_edit.setText(value)

    def upload_image(self):
        # Здесь должен быть код для открытия диалога выбора файла и добавления изображения
        pass

    def show_products_table_page(self):
        # Создаем страницу с таблицей товаров
        products_table_widget = QWidget()
        products_table_layout = QVBoxLayout(products_table_widget)

        # Добавление заголовка и кнопок на одной горизонтали
        page_title_layout = QHBoxLayout()
        label_title = QLabel("Товары")
        label_title.setAlignment(Qt.AlignCenter)
        label_title.setStyleSheet("font-size: 20px; font-weight: bold;")
        page_title_layout.addWidget(label_title)

        show_table_button = QPushButton("Показ таблицы")
        show_table_button.clicked.connect(self.update_table_product)
        show_table_button.setStyleSheet("background-color: transparent; border: none;")
        page_title_layout.addWidget(show_table_button)

        add_product_button = QPushButton("Добавление товара")
        add_product_button.clicked.connect(self.show_products_page)
        add_product_button.setStyleSheet("background-color: transparent; border: none;")
        page_title_layout.addWidget(add_product_button)

        page_title_layout.addStretch(1)
        products_table_layout.addLayout(page_title_layout)

        # Создаем таблицу для отображения товаров
        self.products_table = QTableWidget()
        self.products_table.setColumnCount(7)
        self.products_table.setHorizontalHeaderLabels(
            ["Код товара", "Наименование", "Категория", "Описание", "Цена", "Количество", "Действия"])
        self.products_table.verticalHeader().setVisible(False)

        self.products_table.horizontalHeader().setStretchLastSection(True)
        self.products_table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.products_table.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)

        products_table_layout.addWidget(self.products_table)
        self.stacked_widget.addWidget(products_table_widget)
        self.stacked_widget.setCurrentWidget(products_table_widget)
        self.update_table_product()

    def delete_order(self, order_id):
        # Получение данных удаленного заказа
        cursor = self.connection.cursor()
        cursor.execute(
            f"SELECT Номер_заказа, ФИО_клиента, Код_товара, Количество, Цена, Общая_цена FROM Заказы WHERE Код_заказа = {order_id};")
        order_data = cursor.fetchone()

        # Проверка на наличие данных заказа
        if order_data is None:
            print(f"Заказа с кодом {order_id} не существует!")
            return

        # Генерируем случайный 'Код_заказа'
        new_order_id = random.randint(0, 1000000)

        # Перенос заказа в историю с добавлением нового 'Код_заказа', изменением статуса на "Отменено" и добавлением сегодняшней даты
        cursor.execute(f"""
            INSERT INTO История_заказов(Код_заказа, Номер_заказа, ФИО_клиента, Статус, Дата, Код_товара, Количество, Цена, Общая_цена)
            VALUES ({new_order_id}, '{order_data[0]}', '{order_data[1]}', 'Отменено', '{datetime.now().date().strftime("%Y-%m-%d")}',
            '{order_data[2]}','{order_data[3]}','{order_data[4]}', '{order_data[5]}')
        """)

        # Фактическое удаление заказа
        cursor.execute(f"DELETE FROM Заказы WHERE Код_заказа = {order_id};")

        # Применение изменений
        self.connection.commit()
        cursor.close()

        # Оставаться на той же странице после выполнения
        self.show_products_table_page()
    def update_table_product(self):
        while self.products_table.rowCount() > 0:
            self.products_table.removeRow(0)

        # Забираем данные о товарах из базы данных и заполняем таблицу
        cursor = self.connection.cursor()
        cursor.execute(
            "SELECT Код_товара, Наименование, Категория, Описание, Цена, Количество FROM Товары;")
        rows = cursor.fetchall()

        for row_num, row_data in enumerate(rows):
            if row_data is not None:  # это уберёт строки с None
                self.products_table.insertRow(row_num)
                for column_num, data in enumerate(row_data):
                    item = QTableWidgetItem(str(data))
                    self.products_table.setItem(row_num, column_num, item)

            # Создание кнопок действий
            action_widget = QWidget()
            action_layout = QHBoxLayout(action_widget)

            # Создание кнопки редактирования со значком карандаша
            edit_button = QPushButton()
            edit_icon = QPixmap('icons/edit.png')
            edit_button.setIcon(QIcon(edit_icon))
            edit_button.clicked.connect(lambda checked, row=row_data: self.open_edit_window(row))
            edit_button.setStyleSheet("border: none;")  # Установка стиля без границы
            action_layout.addWidget(edit_button)

            # Создание кнопки удаления со значком мусорки
            delete_button = QPushButton()
            delete_button.setIcon(QIcon("icons/delete.png"))
            delete_button.clicked.connect(lambda: self.delete_product(row_data[0])) # Добавьте это на том же уровне, где создается кнопка удаления
            delete_button.setStyleSheet("border: none;")
            action_layout.addWidget(delete_button)

            action_layout.setContentsMargins(0, 0, 0, 0)
            action_layout.setSpacing(0)
            action_widget.setLayout(action_layout)

            self.products_table.setCellWidget(row_num, len(row_data), action_widget)

        cursor.close()

    def delete_product(self, product_id):
        cursor = self.connection.cursor()

        # Проверяем наличие связанных заказов перед удалением товара
        cursor.execute(
            f"SELECT COUNT(*) FROM Заказы WHERE Код_товара = {product_id};"
        )
        linked_orders_count = cursor.fetchone()[0]
        if linked_orders_count > 0:
            print(f"Невозможно удалить товар, так как на него есть ссылки в таблице 'Заказы'!")
            return

        # Если нет связанных заказов, удаляем товар
        cursor.execute(
            f"DELETE FROM Товары WHERE Код_товара = {product_id};"
        )

        # Сохраняем изменения и закрываем соединение
        self.connection.commit()
        cursor.close()

        print(f"Товар с кодом товара {product_id} успешно удален.")

    def open_edit_window(self, row):
        # Создаем новое окно для редактирования
        while self.stacked_widget.count() > 0:
            widget = self.stacked_widget.widget(0)
            self.stacked_widget.removeWidget(widget)
            widget.deleteLater()

        edit_page_widget = QWidget()
        main_layout = QVBoxLayout(edit_page_widget)

        # Layout для полей ввода и изображения
        fields_layout = QVBoxLayout()
        image_layout = QVBoxLayout()

        # Создаем поля ввода и заполняем их данными из выбранного ряда
        fields = ['Код товара', 'Наименование', 'Категория', 'Описание', 'Цена', 'Количество']
        product_fields = {}

        for name, cell in zip(fields, row):
            field_layout = QHBoxLayout()
            label = QLabel(name)
            line_edit = QLineEdit()
            line_edit.setText(str(cell))
            line_edit.setFixedWidth(200)
            field_layout.addWidget(label)
            field_layout.addWidget(line_edit)
            fields_layout.addLayout(field_layout)
            product_fields[name] = line_edit  # Сохраняем QLineEdit объект

        # Добавляем поле для изображения
        image_label = QLabel("Изображение")
        image_field = QLineEdit()

        # Получение изображения из БД по 'Код_товара'
        cursor = self.connection.cursor()
        cursor.execute("SELECT Изображение FROM Товары WHERE Код_товара = %s", (row[0],))  # row[0] - 'Код_товара'
        image_data = cursor.fetchone()
        if image_data is not None:
            image_field.setText(str(image_data[0]))
        cursor.close()

        image_field.setFixedWidth(200)
        image_layout.addWidget(image_label)
        image_layout.addWidget(image_field)
        product_fields['Изображение'] = image_field  # Сохраняем поле для изображения

        # Объединяем поля для ввода и поле для изображения в одну строку
        field_and_image_layout = QHBoxLayout()
        field_and_image_layout.addLayout(fields_layout)
        field_and_image_layout.addLayout(image_layout)
        main_layout.addLayout(field_and_image_layout)

        save_changes_button = QPushButton('Подтверждение')
        save_changes_button.setStyleSheet("background-color: black; color: white; border: none; padding: 10px;")
        save_changes_button.clicked.connect(lambda: self.update_database(product_fields))  # передаем словарь полей

        main_layout.addWidget(save_changes_button)

        edit_page_widget.setLayout(main_layout)
        self.stacked_widget.addWidget(edit_page_widget)
        self.stacked_widget.setCurrentWidget(edit_page_widget)

    def update_database(self, product_fields):
        # получаем данные из полей ввода
        product_data = {key: line_edit.text() if line_edit.text() else None for key, line_edit in
                        product_fields.items()}

        if not product_data['Изображение']:
            print("Не все обязательные поля заполнены!")
            return

        try:
            cursor = self.connection.cursor()

            data_tuple = (
                product_data['Код товара'],
                product_data['Наименование'],
                product_data['Описание'],
                product_data['Категория'],
                product_data['Цена'],
                int(product_data['Количество']),
                product_data['Изображение'],  # добавлено поле изображения
                product_data['Код товара']  # Код товара, который нужно обновить
            )

            cursor.execute(
                "UPDATE Товары SET Код_товара = %s, Наименование = %s, Описание = %s, Категория = %s, Цена = %s, Количество = %s, Изображение = %s WHERE Код_товара = %s",
                data_tuple)
            self.connection.commit()
            print("Изменения успешно сохранены.")

            # Если все успешно, переходим на страницу со списком товаров
            self.show_products_table_page()

        except Exception as e:
            print(f"Произошла ошибка при обновлении данных в БД: {e}")

        finally:
            cursor.close()

    def make_delete_order_lambda(self, order_id):
        return lambda: self.delete_order(order_id)

    def show_clients_table_page(self):
        # Создаем страницу с таблицей клиентов
        clients_table_widget = QWidget()
        clients_table_layout = QVBoxLayout(clients_table_widget)

        # Добавление заголовка и кнопок на одной горизонтали
        page_title_layout = QHBoxLayout()
        label_title = QLabel("Клиенты")
        label_title.setAlignment(Qt.AlignCenter)
        label_title.setStyleSheet("font-size: 20px; font-weight: bold;")
        page_title_layout.addWidget(label_title)

        show_table_button = QPushButton("Показ таблицы")
        show_table_button.clicked.connect(self.update_table_client)
        show_table_button.setStyleSheet("background-color: transparent; border: none;")
        page_title_layout.addWidget(show_table_button)

        page_title_layout.addStretch(1)
        clients_table_layout.addLayout(page_title_layout)

        # Создаем таблицу для отображения клиентов
        self.clients_table = QTableWidget()
        self.clients_table.setColumnCount(2)
        self.clients_table.setHorizontalHeaderLabels(
            ["Код клиента", "ФИО",])
        self.clients_table.verticalHeader().setVisible(False)

        self.clients_table.horizontalHeader().setStretchLastSection(True)
        self.clients_table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.clients_table.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)

        clients_table_layout.addWidget(self.clients_table)
        self.stacked_widget.addWidget(clients_table_widget)
        self.stacked_widget.setCurrentWidget(clients_table_widget)
        self.update_table_client()  # Вызов функции для отображения данных в таблице

    def update_table_client(self):
        # Версия функции с print()
        print("Окно Клиентов")  # Вывод в терминал при нажатии кнопки

        # Удалить текущие строки в таблице
        while self.clients_table.rowCount() > 0:
            self.clients_table.removeRow(0)

        # Забираем данные о клиентах из базы данных и заполняем таблицу
        cursor = self.connection.cursor()
        cursor.execute(
            "SELECT Код_клиента, ФИО, Телеграм_id FROM Клиенты;")
        rows = cursor.fetchall()
        for row_num, row_data in enumerate(rows):
            self.clients_table.insertRow(row_num)
            for column_num, data in enumerate(row_data):
                item = QTableWidgetItem(str(data))
                self.clients_table.setItem(row_num, column_num, item)


        cursor.close()


    def show_history_orders(self):
        # Создаем страницу с таблицей истории заказов
        history_orders_widget = QWidget()
        history_orders_layout = QVBoxLayout(history_orders_widget)

        # Добавление заголовка и кнопок на одной горизонтали
        page_title_layout = QHBoxLayout()
        label_title = QLabel("История заказов")
        label_title.setAlignment(Qt.AlignCenter)
        label_title.setStyleSheet("font-size: 20px; font-weight: bold;")
        page_title_layout.addWidget(label_title)

        show_table_button = QPushButton("Показ таблицы")
        show_table_button.clicked.connect(self.update_table)
        show_table_button.setStyleSheet("background-color: transparent; border: none;")
        page_title_layout.addWidget(show_table_button)

        page_title_layout.addStretch(1)
        history_orders_layout.addLayout(page_title_layout)

        # Создание и настройка таблицы истории заказов
        self.history_table = QTableWidget()
        self.history_table.setColumnCount(6)
        self.history_table.setHorizontalHeaderLabels(
            ["Код заказа", "Номер заказа", "ФИО клиента", "Статус", "Дата и время","Действия"])
        self.history_table.verticalHeader().setVisible(False)

        self.history_table.horizontalHeader().setStretchLastSection(True)
        self.history_table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.history_table.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
        history_orders_layout.addWidget(self.history_table)
        self.stacked_widget.addWidget(history_orders_widget)
        self.stacked_widget.setCurrentWidget(history_orders_widget)
        self.update_table()

    def update_table(self):
        # Удалить текущие строки в таблице
        while self.history_table.rowCount() > 0:
            self.history_table.removeRow(0)

        # Заполнение таблицы данными из базы данных
        cursor = self.connection.cursor()
        cursor.execute(
            "SELECT Код_заказа, Номер_заказа, ФИО_клиента, Статус, Дата FROM История_заказов;")
        rows = cursor.fetchall()
        for row_num, row_data in enumerate(rows):
            self.history_table.insertRow(row_num)
            for column_num, data in enumerate(row_data):
                item = QTableWidgetItem(str(data))
                self.history_table.setItem(row_num, column_num, item)

            # Создание кнопок действий
            action_widget = QWidget()
            action_layout = QHBoxLayout(action_widget)

            # Создание кнопки редактирования со значком глаза
            view_button = QPushButton()
            view_icon = QPixmap('icons/eye.png')
            view_button.setIcon(QIcon(view_icon))
            # Связывание кнопки с кодом заказа
            view_button.clicked.connect(lambda checked, row=row_data: self.edit_product_1(row[0]))
            view_button.setStyleSheet("border: none;")
            action_layout.addWidget(view_button)

            action_layout.setContentsMargins(0, 0, 0, 0)
            action_layout.setSpacing(0)
            action_widget.setLayout(action_layout)

            self.history_table.setCellWidget(row_num, len(row_data), action_widget)

        cursor.close()

    def edit_product_1(self, order_id):
        cursor = self.connection.cursor()
        cursor.execute(f"SELECT Статус FROM История_заказов WHERE Код_заказа = {order_id};")
        order_status = cursor.fetchone()

        if order_status is not None and order_status[0] == "Отменено":
            print("Запрет на просмотр.")
            cursor.close()
            return

        # Удаление текущих строк в таблице
        while self.history_table.rowCount() > 0:
            self.history_table.removeRow(0)

        detail_widget = QWidget()
        detail_layout = QVBoxLayout(detail_widget)

        # Добавление заголовка и кнопок на одной горизонтали
        page_title_layout = QHBoxLayout()
        label_title = QLabel("История заказов")
        label_title.setAlignment(Qt.AlignCenter)
        label_title.setStyleSheet("font-size: 20px; font-weight: bold;")
        page_title_layout.addWidget(label_title)

        show_table_button = QPushButton("Показ таблицы")
        show_table_button.clicked.connect(self.show_history_orders)
        show_table_button.setStyleSheet("background-color: transparent; border: none;")
        page_title_layout.addWidget(show_table_button)

        view_order_button = QPushButton("Просмотр заказа")
        view_order_button.setStyleSheet("background-color: transparent; border: none;")
        page_title_layout.addWidget(view_order_button)

        detail_layout.addLayout(page_title_layout)

        # Заголовки таблицы деталей товара
        detail_table = QTableWidget()
        detail_table.setColumnCount(4)
        detail_table.setHorizontalHeaderLabels(
            ["Наименование товара", "Кол-во", "Цена", "Общая цена"])
        detail_table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)  # Расширение таблицы

        # Заполнение таблицы данными из базы данных
        cursor.execute(
            f"SELECT Код_товара, Количество, Цена, Общая_цена FROM История_заказов WHERE Код_заказа = {order_id};")
        rows = cursor.fetchall()
        print(rows)
        for row_num, row_data in enumerate(rows):
            detail_table.insertRow(row_num)

            # Получение имени товара
            product_id = row_data[0]
            print("Впфмц",product_id)
            cursor.execute(f"SELECT Наименование FROM Товары WHERE Код_товара = {product_id};")
            product_name = cursor.fetchone()

            if product_name is not None:
                product_name = product_name[0]
            else:
                product_name = "Неопределенный товар"

            # Замена кода товара на наименование товара
            row_data = list(row_data)
            row_data[0] = product_name

            for column_num, data in enumerate(row_data):
                item = QTableWidgetItem(str(data))
                detail_table.setItem(row_num, column_num, item)

        detail_layout.addWidget(detail_table)
        self.stacked_widget.addWidget(detail_widget)
        self.stacked_widget.setCurrentWidget(detail_widget)
        cursor.close()

    def show_employees_table_page(self):
        # Создаем страницу с таблицей сотрудников
        employees_table_widget = QWidget()
        employees_table_layout = QVBoxLayout(employees_table_widget)

        # Добавление заголовка и кнопок на одной горизонтали
        page_title_layout = QHBoxLayout()
        label_title = QLabel("Сотрудники")
        label_title.setAlignment(Qt.AlignCenter)
        label_title.setStyleSheet("font-size: 20px; font-weight: bold;")
        page_title_layout.addWidget(label_title)

        show_table_button = QPushButton("Показ таблицы")
        show_table_button.clicked.connect(self.update_table_employees)
        show_table_button.setStyleSheet("background-color: transparent; border: none;")
        page_title_layout.addWidget(show_table_button)

        add_employee_button = QPushButton("Добавление сотрудника")
        add_employee_button.clicked.connect(
            self.show_add_employee_page)  # Предположительно у вас есть функция для добавления сотрудника
        add_employee_button.setStyleSheet("background-color: transparent; border: none;")
        page_title_layout.addWidget(add_employee_button)

        page_title_layout.addStretch(1)
        employees_table_layout.addLayout(page_title_layout)

        # Создаем таблицу для отображения сотрудников
        self.employees_table = QTableWidget()
        self.employees_table.setColumnCount(6)
        self.employees_table.setHorizontalHeaderLabels(
            ["Код сотрудника", "ФИО", "Телефон", "Адрес", "Дата рождения", "Действия"])
        self.employees_table.verticalHeader().setVisible(False)

        self.employees_table.horizontalHeader().setStretchLastSection(True)
        self.employees_table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.employees_table.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)

        employees_table_layout.addWidget(self.employees_table)
        self.stacked_widget.addWidget(employees_table_widget)
        self.stacked_widget.setCurrentWidget(employees_table_widget)
        self.update_table_employees()  # Вызов функции для отображения данных в таблице

    def update_table_employees(self):
        # Удалить текущие строки в таблице
        while self.employees_table.rowCount() > 0:
            self.employees_table.removeRow(0)

        # Забираем данные о сотрудниках из базы данных и заполняем таблицу
        cursor = self.connection.cursor()
        cursor.execute(
            "SELECT Код_сотрудника, ФИО, Телефон, Адрес, Дата_рождения FROM Сотрудники;")
        rows = cursor.fetchall()
        for row_num, row_data in enumerate(rows):
            self.employees_table.insertRow(row_num)
            for column_num, data in enumerate(row_data):
                item = QTableWidgetItem(str(data))
                self.employees_table.setItem(row_num, column_num, item)

            # Создание кнопок действий
            action_widget = QWidget()
            action_layout = QHBoxLayout(action_widget)

            # Создание кнопки редактирования со значком карандаша
            edit_button = QPushButton()
            edit_icon = QPixmap('icons/edit.png')
            edit_button.setIcon(QIcon(edit_icon))
            edit_button.clicked.connect(lambda checked, row=row_data: self.open_edit_employee_window(row))  # Нужно написать правильный метод
            edit_button.setStyleSheet("border: none;")  # Установка стиля без границы
            action_layout.addWidget(edit_button)

            # Создание кнопки удаления со значком мусорки
            delete_button = QPushButton()
            delete_button.setIcon(QIcon("icons/delete.png"))
            delete_button.clicked.connect(lambda: self.delete_employees(row_data[0]))  # Добавьте это на том же уровне, где создается кнопка удаления
            delete_button.setStyleSheet("border: none;")
            action_layout.addWidget(delete_button)

            action_layout.setContentsMargins(0, 0, 0, 0)
            action_layout.setSpacing(0)
            action_widget.setLayout(action_layout)

            self.employees_table.setCellWidget(row_num, len(row_data), action_widget)  # Исправлено здесь

        cursor.close()
    def delete_employees(self, order_id):
        cursor = self.connection.cursor()
        cursor.execute(f"DELETE FROM Сотрудники WHERE Код_сотрудника = {order_id};")
        self.connection.commit()
        cursor.close()

    def open_edit_employee_window(self, row):
        while self.stacked_widget.count() > 0:
            widget = self.stacked_widget.widget(0)
            self.stacked_widget.removeWidget(widget)
            widget.deleteLater()

        edit_page_widget = QWidget()
        main_layout = QVBoxLayout()
        main_layout.setAlignment(Qt.AlignCenter)

        page_title_layout = QHBoxLayout()
        label_title = QLabel("Сотрудники")
        label_title.setAlignment(Qt.AlignCenter)
        label_title.setStyleSheet("font-size: 20px; font-weight: bold;")
        page_title_layout.addWidget(label_title)

        show_table_button = QPushButton("Показ таблицы")
        show_table_button.clicked.connect(self.show_employees_table_page)
        show_table_button.setStyleSheet("background-color: transparent; border: none;")
        page_title_layout.addWidget(show_table_button)

        add_employee_button = QPushButton("Добавление сотрудника")
        add_employee_button.setStyleSheet("background-color: transparent; border: none;")
        page_title_layout.addWidget(add_employee_button)

        main_layout.addLayout(page_title_layout)

        fields = ['Код сотрудника', 'ФИО', 'Телефон', 'Адрес', 'Дата рождения']
        employee_fields = {}

        for name, cell in zip(fields, row):
            field_layout = QVBoxLayout()
            label = QLabel(name)
            label.setAlignment(Qt.AlignCenter)
            line_edit = QLineEdit()
            line_edit.setText(str(cell))
            line_edit.setAlignment(Qt.AlignCenter)
            line_edit.setFixedWidth(300)
            field_layout.addWidget(label)
            field_layout.addWidget(line_edit)
            main_layout.addLayout(field_layout)
            employee_fields[name] = line_edit

        save_changes_button = QPushButton('Подтверждение')
        save_changes_button.setStyleSheet("background-color: black; color: white; border: none; padding: 10px;")
        save_changes_button.setFixedSize(150, 40)
        save_changes_button.clicked.connect(lambda: self.update_database_employee(employee_fields))
        main_layout.addWidget(save_changes_button, alignment=Qt.AlignCenter)

        edit_page_widget.setLayout(main_layout)
        self.stacked_widget.addWidget(edit_page_widget)

    def update_database_employee(self, employee_fields):
        # получаем данные из полей ввода
        employee_data = {key: line_edit.text() for key, line_edit in employee_fields.items()}

        if any(not value for value in employee_data.values()):
            print("Не все обязательные поля заполнены!")
            return

        try:
            cursor = self.connection.cursor()

            data_tuple = (
                employee_data['ФИО'],
                employee_data['Телефон'],
                employee_data['Адрес'],
                employee_data['Дата рождения'],
                employee_data['Код сотрудника']
            )

            cursor.execute(
                "UPDATE Сотрудники SET ФИО = %s, Телефон = %s, Адрес = %s, Дата_рождения = %s WHERE Код_сотрудника = %s",
                data_tuple)
            self.connection.commit()
            print("Изменения успешно сохранены.")

            # Если все успешно, то можно перейти на страницу со списком сотрудников
            self.show_employee_table_page()

        except Exception as e:
            print(f"Произошла ошибка при обновлении данных в БД: {e}")

        finally:
            cursor.close()

    def show_add_employee_page(self):
        # Удаление всех виджетов из QStackedWidget
        while self.stacked_widget.count() > 0:
            widget = self.stacked_widget.widget(0)
            self.stacked_widget.removeWidget(widget)
            widget.deleteLater()

        add_employee_page_widget = QWidget()
        main_layout = QVBoxLayout(add_employee_page_widget)
        main_layout.setAlignment(Qt.AlignCenter)  # Установка выравнивания по центру для основного layout

        # Название страницы
        title_layout = QHBoxLayout()
        title_label = QLabel("Сотрудники")
        title_label.setStyleSheet("font-size: 20px; font-weight: bold;")
        title_label.setAlignment(Qt.AlignCenter)  # Установка выравнивания по центру для метки заголовка
        title_layout.addWidget(title_label)
        main_layout.addLayout(title_layout)

        show_table_button = QPushButton("Показ таблицы")
        show_table_button.clicked.connect(self.show_employees_table_page)
        show_table_button.setStyleSheet("background-color: transparent; border: none;")
        title_layout.addWidget(show_table_button)

        add_employee_button = QPushButton("Добавление сотрудника")
        add_employee_button.setStyleSheet("background-color: transparent; border: none;")
        title_layout.addWidget(add_employee_button)

        # Создание разметки для полей ввода
        fields_layout = QVBoxLayout()
        fields = ['ФИО', 'Телефон', 'Адрес', 'Дата рождения']
        self.employee_fields = {}
        for field in fields:
            field_layout = QVBoxLayout()  # Используем QVBoxLayout здесь
            field_layout.setAlignment(Qt.AlignCenter)  # Установка выравнивания по центру для field_layout
            label = QLabel(field + ":")
            line_edit = QLineEdit()
            line_edit.setAlignment(Qt.AlignCenter)  # Установка выравнивания по центру для поля ввода
            line_edit.setFixedWidth(300)  # Установка фиксированной ширины для поля ввода
            self.employee_fields[field] = line_edit
            field_layout.addWidget(label)
            field_layout.addWidget(line_edit)
            fields_layout.addLayout(field_layout)

        main_layout.addLayout(fields_layout)

        # Кнопка подтверждения
        confirm_button_layout = QHBoxLayout()
        confirm_button_layout.setAlignment(Qt.AlignCenter)  # Установка выравнивания по центру для confirm_button_layout
        confirm_button = QPushButton('Добавление')
        confirm_button.setStyleSheet("background-color: black; color: white; border: none; padding: 10px;")
        confirm_button.setFixedSize(150, 40)  # Установка фиксированного размера для кнопки
        confirm_button.clicked.connect(self.confirm_employee)
        confirm_button_layout.addWidget(confirm_button)
        main_layout.addLayout(confirm_button_layout)

        add_employee_page_widget.setLayout(main_layout)
        self.stacked_widget.addWidget(add_employee_page_widget)
        self.stacked_widget.setCurrentWidget(add_employee_page_widget)

    def confirm_employee(self):
        # Сбор данных из полей ввода
        employee_data = {field: self.employee_fields[field].text().strip() for field in self.employee_fields}

        # Проверяем обязательные поля на заполненность
        if any(not value for value in employee_data.values()):
            print("Не все поля заполнены!")
            return

        # Здесь следует добавить логику для сохранения данных в базу данных
        try:
            cursor = self.connection.cursor()
            cursor.execute(
                "INSERT INTO Сотрудники (ФИО, Телефон, Адрес, Дата_рождения) VALUES (%s, %s, %s, %s)",
                (employee_data['ФИО'], employee_data['Телефон'], employee_data['Адрес'], employee_data['Дата рождения'])
            )
            self.connection.commit()
            print("Сотрудник успешно добавлен.")
        except Exception as e:
            print(f"Ошибка при добавлении сотрудника: {e}")
        finally:
            cursor.close()

    def show_report_page(self):
        report_widget = QWidget()
        report_layout = QVBoxLayout(report_widget)

        label_title = QLabel("Отчет")
        label_title.setAlignment(Qt.AlignTop | Qt.AlignLeft)
        label_title.setStyleSheet("font-size: 20px; font-weight: bold; margin: 10px 0px 0px 10px;")
        report_layout.addWidget(label_title)

        period_layout = QVBoxLayout()

        self.report_type_combo = QComboBox()
        self.report_type_combo.addItems([
            "Отчет по истории заказов",
            "Отчет по количеству товаров",
            "Отчет по сотрудникам",
            "Отчет по прибыли",
            "Отчет по клиентам"
        ])
        period_layout.addWidget(self.report_type_combo)

        period_label = QLabel("Период:")
        period_label.setAlignment(Qt.AlignLeft)
        period_label.setStyleSheet("font-size: 14px; font-weight: bold; margin-left: 10px; margin-bottom: 5px;")
        period_layout.addWidget(period_label)

        date_input_layout = QHBoxLayout()
        self.start_date_edit = QLineEdit()
        self.start_date_edit.setPlaceholderText("Начальная дата (ДД.ММ.ГГГГ)")
        date_input_layout.addWidget(self.start_date_edit)
        self.end_date_edit = QLineEdit()
        self.end_date_edit.setPlaceholderText("Конечная дата (ДД.ММ.ГГГГ)")
        date_input_layout.addWidget(self.end_date_edit)
        period_layout.addLayout(date_input_layout)

        report_layout.addLayout(period_layout)

        confirm_button = QPushButton("Подтверждение")
        confirm_button.setStyleSheet(
            "background-color: black; color: white; border: none; padding: 10px; margin-top: 10px;")
        confirm_button.clicked.connect(self.on_confirm_report)
        report_layout.addWidget(confirm_button, alignment=Qt.AlignCenter)

        report_widget.setLayout(report_layout)
        self.stacked_widget.addWidget(report_widget)
        self.stacked_widget.setCurrentWidget(report_widget)

    def on_confirm_report(self):
        start_date = self.start_date_edit.text().strip()
        end_date = self.end_date_edit.text().strip()

        report_type_to_query = {
            "Отчет по истории заказов": "SELECT * FROM История_заказов WHERE Дата BETWEEN %s AND %s;",
            "Отчет по количеству товаров": "SELECT Код_товара, Наименование,Категория, Количество FROM Товары",
            "Отчет по сотрудникам": "SELECT * FROM Сотрудники",
            "Отчет по прибыли": "SELECT SUM(Общая_цена) FROM История_заказов WHERE Дата BETWEEN %s AND %s;",
            "Отчет по клиентам": "SELECT * FROM Клиенты",
        }

        selected_report_type = self.report_type_combo.currentText()
        selected_query = report_type_to_query.get(selected_report_type)

        headers = ["Код_заказа", "Номер_заказа", "ФИО_клиента", "Статус", "Дата_и_время", "Код_товара",
                   "Количество", "Цена", "Общая_цена"]

        if selected_report_type == "Отчет по количеству товаров":
            headers = ["Код Товара", "Наименование", "Категория", "Количество"]
        elif selected_report_type == "Отчет по сотрудникам":
            headers = ["Код_сотрудника", "ФИО", "Телефон", "Адрес", "дата рождения"]  # Пример заголовков колонок
        elif selected_report_type == "Отчет по прибыли":
            headers = ["Общая прибыль"]
        elif selected_report_type == "Отчет по клиентам":
            headers = ["ФИО", "Телефон", "Адрес"]  # Пример заголовков колонок

        doc = Document()

        header = doc.add_heading(level=1)
        run = header.add_run(selected_report_type)
        run.bold = True
        run.font.size = Pt(14)
        header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'

        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = header

        cursor = self.connection.cursor()
        if selected_report_type in ["Отчет по истории заказов", "Отчет по прибыли"]:
            cursor.execute(selected_query, (start_date, end_date))
        else:
            cursor.execute(selected_query)

        orders_data = cursor.fetchall()

        for row_data in orders_data:
            cells = table.add_row().cells
            for i, value in enumerate(row_data):
                cells[i].text = str(value)

        cursor.close()

        doc.save('Отчет.docx')
        print("Word отчет успешно создан.")
def main():
    app = QApplication(sys.argv)
    window = ComputerStoreApp()
    window.show()
    sys.exit(app.exec_())


if __name__ == "__main__":
    main()